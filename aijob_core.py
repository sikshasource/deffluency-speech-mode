# import re
# from collections import Counter
# from typing import List, Dict, Any, Optional

# try:
#     import pdfplumber
# except ImportError:
#     pdfplumber = None

# try:
#     import docx
# except ImportError:
#     docx = None

# STOPWORDS = {
#     "the", "and", "for", "with", "you", "your", "a", "an", "to", "of", "in",
#     "on", "at", "by", "as", "is", "are", "be", "this", "that", "will", "our",
#     "we", "they", "their", "from", "or", "it", "have", "has", "had", "i",
#     "about", "us", "who", "what", "when", "where", "how", "why"
# }

# def tokenize(text: str) -> List[str]:
#     text = text.lower()
#     text = re.sub(r"[^a-z0-9+]", " ", text)
#     return [t for t in text.split() if t and t not in STOPWORDS]

# def extract_text_from_pdf(file_path: str) -> str:
#     if not pdfplumber:
#         raise ImportError("pdfplumber library not installed. Install with 'pip install pdfplumber'")
#     text = []
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             text.append(page.extract_text() or "")
#     return "\n".join(text)

# def extract_text_from_docx(file_path: str) -> str:
#     if not docx:
#         raise ImportError("python-docx library not installed. Install with 'pip install python-docx'")
#     document = docx.Document(file_path)
#     return "\n".join([p.text for p in document.paragraphs])

# def extract_resume_text(file_path: str) -> str:
#     file_path_lower = file_path.lower()
#     if file_path_lower.endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path_lower.endswith(".docx"):
#         return extract_text_from_docx(file_path)
#     else:
#         with open(file_path, "r", encoding="utf-8") as f:
#             return f.read()

# def extract_role(job_description: str) -> str:
#     lines = [l.strip() for l in job_description.splitlines() if l.strip()]
#     return lines[0][:120] if lines else "insufficient data"

# def extract_experience_required(job_description: str) -> str:
#     match = re.search(r"(\d+)\+?\s+years", job_description, flags=re.IGNORECASE)
#     if match:
#         return f"{match.group(1)}+ years"
#     range_match = re.search(r"(\d+)\s*-\s*(\d+)\s+years", job_description, flags=re.IGNORECASE)
#     if range_match:
#         return f"{range_match.group(1)}-{range_match.group(2)} years"
#     return "not specified"

# def extract_skills_from_section(text: str) -> List[str]:
#     skills: List[str] = []
#     lines = text.splitlines()
#     headers = ("requirements", "responsibilities", "qualifications", "skills", "what you will do")
#     current_section = None

#     for line in lines:
#         stripped = line.strip()
#         lower = stripped.lower()

#         if any(h in lower for h in headers):
#             current_section = "skills"
#             continue

#         if current_section == "skills":
#             if re.match(r"^[A-Z][A-Za-z\s]{0,40}:?$", stripped) and len(stripped.split()) <= 6:
#                 current_section = None
#                 continue

#             if stripped.startswith(("-", "*", "•", "·")):
#                 item = stripped.lstrip("-*•· ").strip()
#                 if item:
#                     skills.append(item)
#             elif stripped and len(stripped) <= 80:
#                 skills.append(stripped)

#     seen = set()
#     unique = []
#     for s in skills:
#         key = s.lower()
#         if key not in seen:
#             seen.add(key)
#             unique.append(s)
#     return unique

# def extract_keywords(job_description: str, max_keywords: int = 15) -> List[str]:
#     tokens = tokenize(job_description)
#     counts = Counter(tokens)
#     common = [w for w, _ in counts.most_common() if len(w) > 3]
#     return common[:max_keywords]

# def extract_hidden_requirements(job_description: str) -> List[str]:
#     text_lower = job_description.lower()
#     patterns = {
#         "fast-paced": "fast-paced environment",
#         "self-starter": "self-starter",
#         "work independently": "work independently",
#         "cross-functional": "cross-functional collaboration",
#         "stakeholders": "stakeholder management",
#         "deadline": "ability to meet deadlines",
#         "ambiguity": "comfort with ambiguity",
#         "ownership": "sense of ownership",
#         "leadership": "leadership responsibilities",
#         "manage team": "team management"
#     }
#     hidden = [phrase for key, phrase in patterns.items() if key in text_lower]
#     return list(dict.fromkeys(hidden))

# def extract_resume_skills(resume: str) -> List[str]:
#     skills: List[str] = []
#     lines = resume.splitlines()
#     in_skills_section = False

#     for line in lines:
#         stripped = line.strip()
#         lower = stripped.lower()

#         if "skills" in lower and len(lower) <= 40:
#             in_skills_section = True
#             continue

#         if in_skills_section:
#             if re.match(r"^[A-Z][A-Za-z\s]{0,40}:?$", stripped) and "skills" not in lower:
#                 in_skills_section = False
#             else:
#                 if stripped.startswith(("-", "*", "•", "·")):
#                     items = stripped.lstrip("-*•· ").split(",")
#                     skills.extend([i.strip() for i in items if i.strip()])
#                 elif "," in stripped:
#                     skills.extend([i.strip() for i in stripped.split(",") if i.strip()])
#                 elif stripped:
#                     skills.append(stripped)
#                 continue

#         if stripped.startswith(("-", "*", "•", "·")):
#             skills.append(stripped.lstrip("-*•· ").strip())

#     seen = set()
#     unique = []
#     for s in skills:
#         key = s.lower()
#         if key not in seen:
#             seen.add(key)
#             unique.append(s)
#     return unique

# def get_overlap_keywords(job_keywords: List[str], resume_text: str) -> List[str]:
#     resume_tokens = set(tokenize(resume_text))
#     return [k for k in job_keywords if k.lower() in resume_tokens]

# def analyze_resume_strengths(
#     job_skills: List[str],
#     resume_skills: List[str],
#     resume_text: str
# ) -> Dict[str, List[str]]:
#     strengths, missing = [], []
#     resume_lower = [s.lower() for s in resume_skills]

#     for js in job_skills:
#         if js.lower() in resume_lower:
#             strengths.append(js)
#         else:
#             missing.append(js)

#     weak_areas: List[str] = []

#     if not re.search(r"\d+\+?\s+years", resume_text, flags=re.IGNORECASE):
#         weak_areas.append("Years of experience not clearly specified")

#     if job_skills and len(strengths) / len(job_skills) < 0.5:
#         weak_areas.append("Significant gap between job-required skills and listed skills")

#     return {"strengths": strengths, "missing": missing, "weak_areas": weak_areas}

# def compute_match_score(
#     job_skills: List[str],
#     resume_skills: List[str],
#     job_keywords: List[str],
#     resume_text: str
# ) -> int:
#     if job_skills:
#         skill_overlap = sum(1 for js in job_skills if js.lower() in [s.lower() for s in resume_skills])
#         skill_score = (skill_overlap / len(job_skills)) * 70.0
#     else:
#         skill_score = 0.0

#     if job_keywords:
#         keyword_overlap = len(get_overlap_keywords(job_keywords, resume_text))
#         keyword_score = (keyword_overlap / len(job_keywords)) * 30.0
#     else:
#         keyword_score = 0.0

#     return max(0, min(100, int(round(skill_score + keyword_score))))

# def decide_application(match_score: int) -> str:
#     if match_score >= 70:
#         return "APPLY"
#     if match_score >= 50:
#         return "MAYBE"
#     return "REJECT"

# def confidence_level(match_score: int, job_skills: List[str], resume_skills: List[str]) -> str:
#     if not job_skills:
#         return "LOW"
#     overlap = sum(1 for js in job_skills if js.lower() in [s.lower() for s in resume_skills])
#     if match_score >= 70 and overlap >= max(1, len(job_skills) // 2):
#         return "HIGH"
#     if 50 <= match_score < 70 and overlap > 0:
#         return "MEDIUM"
#     return "LOW"

# def build_optimized_resume(candidate_resume: str, job_keywords: List[str], resume_skills: List[str]) -> str:
#     if not candidate_resume.strip():
#         return "insufficient data"

#     from_text_tokens = set(tokenize(candidate_resume))
#     kept_keywords = [k for k in job_keywords if k.lower() in from_text_tokens]
#     core_skills = [s for s in resume_skills if len(s) <= 60][:15]

#     summary_lines = []
#     if core_skills:
#         summary_lines.append("SUMMARY")
#         summary_lines.append("Experienced professional with strengths in: " + ", ".join(core_skills) + ".")
#     if kept_keywords:
#         summary_lines.append("")
#         summary_lines.append("KEYWORDS")
#         summary_lines.append(", ".join(sorted(set(kept_keywords))))

#     if summary_lines:
#         return "\n".join(summary_lines) + "\n\n---\n\n" + candidate_resume.strip()
#     return candidate_resume.strip()

# def build_cover_letter(
#     job_role: str,
#     job_keywords: List[str],
#     strengths: List[str],
#     candidate_resume: str
# ) -> str:
#     if not candidate_resume.strip():
#         return "insufficient data"

#     role = job_role if job_role != "insufficient data" else "the role"
#     key_strengths = strengths[:5]
#     resume_tokens = set(tokenize(candidate_resume))
#     aligned_keywords = [k for k in job_keywords if k.lower() in resume_tokens][:5]

#     strengths_text = ", ".join(key_strengths) if key_strengths else "the skills outlined in my resume"
#     keywords_text = ", ".join(aligned_keywords) if aligned_keywords else ""

#     p1 = (
#         f"I am writing to express my interest in {role}. "
#         f"After reviewing the job description, I believe my background and skills are well aligned with your needs."
#     )
#     if keywords_text:
#         p2 = (
#             f"Throughout my experience, I have developed strong capabilities in {keywords_text}. "
#             f"These areas are reflected in my work, where I have applied them to deliver meaningful results."
#         )
#     else:
#         p2 = (
#             "My experience spans multiple responsibilities where I have consistently focused on delivering "
#             "high-quality outcomes and collaborating effectively with stakeholders."
#         )
#     p3 = (
#         f"My key strengths include {strengths_text}. "
#         "I am comfortable taking ownership of my work, communicating clearly, and adapting to changing priorities."
#     )
#     p4 = (
#         "I would welcome the opportunity to further discuss how my experience can contribute to your team’s goals. "
#         "Thank you for considering my application."
#     )
#     return "\n\n".join([p1, p2, p3, p4])

# def analyze_application(
#     job_description: str,
#     candidate_resume_text: str,
#     extra_context: Optional[str] = None
# ) -> Dict[str, Any]:
#     job_description = job_description or ""
#     candidate_resume_text = candidate_resume_text or ""

#     role = extract_role(job_description) if job_description.strip() else "insufficient data"
#     experience_required = extract_experience_required(job_description) if job_description.strip() else "insufficient data"
#     job_skills = extract_skills_from_section(job_description) if job_description.strip() else []
#     keywords = extract_keywords(job_description) if job_description.strip() else []
#     hidden = extract_hidden_requirements(job_description) if job_description.strip() else []

#     job_analysis = {
#         "role": role,
#         "skills": job_skills,
#         "experience_required": experience_required,
#         "keywords": keywords or ["insufficient data"],
#         "hidden_requirements": hidden or ["insufficient data"]
#     }

#     resume_skills = extract_resume_skills(candidate_resume_text) if candidate_resume_text.strip() else []
#     if candidate_resume_text.strip():
#         strength_analysis = analyze_resume_strengths(job_skills, resume_skills, candidate_resume_text)
#         resume_strengths = strength_analysis["strengths"]
#         missing_skills = strength_analysis["missing"]
#         weak_areas = strength_analysis["weak_areas"]
#     else:
#         resume_strengths = []
#         missing_skills = job_skills
#         weak_areas = ["Candidate resume not provided or empty"]

#     resume_analysis = {
#         "strengths": resume_strengths,
#         "missing_skills": missing_skills,
#         "weak_areas": weak_areas
#     }

#     score = compute_match_score(job_skills, resume_skills, keywords, candidate_resume_text)
#     decision = decide_application(score)
#     conf = confidence_level(score, job_skills, resume_skills)

#     if decision == "APPLY":
#         reason = f"Match score {score}: substantial overlap; suitable to submit application."
#     elif decision == "MAYBE":
#         reason = f"Match score {score}: partial overlap; may be worthwhile."
#     else:
#         reason = f"Match score {score}: limited overlap; unlikely to be competitive."

#     optimized_resume = build_optimized_resume(candidate_resume_text, keywords, resume_skills)
#     cover_letter = build_cover_letter(role, keywords, resume_strengths, candidate_resume_text)

#     return {
#         "job_analysis": job_analysis,
#         "resume_analysis": resume_analysis,
#         "match_score": score,
#         "decision": decision,
#         "reason": reason,
#         "confidence": conf,
#         "optimized_resume": optimized_resume,
#         "cover_letter": cover_letter
#     }
# import os
# import json
# from flask import Flask, render_template_string, request, redirect, url_for, flash
# from werkzeug.utils import secure_filename

# from aijob_core import analyze_application, extract_resume_text

# app = Flask(__name__)
# app.secret_key = "change-this-secret-key"

# UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "uploads")
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# # Single-page template with a modern Bootstrap layout
# TEMPLATE = """
# <!doctype html>
# <html lang="en">
# <head>
#   <meta charset="utf-8">
#   <title>AI Job Match Assistant</title>
#   <meta name="viewport" content="width=device-width, initial-scale=1">
#   <link
#     href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css"
#     rel="stylesheet"
#   >
#   <style>
#     body {
#       background: #0b1727;
#       color: #f0f4ff;
#       font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
#     }
#     .navbar-brand {
#       font-weight: 700;
#       letter-spacing: 0.05em;
#     }
#     .card {
#       border-radius: 14px;
#       border: 1px solid rgba(255,255,255,0.06);
#       background: radial-gradient(circle at top left, #1f2a3c, #111827);
#       box-shadow: 0 18px 45px rgba(0,0,0,0.55);
#     }
#     .card-header {
#       border-bottom: 1px solid rgba(255,255,255,0.06);
#       background: transparent;
#       font-weight: 600;
#       letter-spacing: 0.04em;
#       text-transform: uppercase;
#       font-size: 0.8rem;
#       color: #9ca3af;
#     }
#     .pill {
#       display: inline-flex;
#       align-items: center;
#       padding: 0.2rem 0.7rem;
#       border-radius: 999px;
#       font-size: 0.75rem;
#       border: 1px solid rgba(255,255,255,0.15);
#       margin: 0.1rem;
#       background: rgba(15,23,42,0.85);
#     }
#     .score-circle {
#       width: 110px;
#       height: 110px;
#       border-radius: 50%;
#       display: flex;
#       align-items: center;
#       justify-content: center;
#       font-size: 1.6rem;
#       font-weight: 700;
#       margin: 0 auto;
#       color: #0b1727;
#     }
#     .score-low { background: linear-gradient(135deg, #fecaca, #f97373); }
#     .score-medium { background: linear-gradient(135deg, #fee2b3, #f59e0b); }
#     .score-high { background: linear-gradient(135deg, #bbf7d0, #22c55e); }
#     textarea {
#       background: #020617;
#       border-radius: 10px;
#       color: #e5e7eb;
#       border: 1px solid rgba(148,163,184,0.6);
#       font-size: 0.9rem;
#     }
#     textarea:focus {
#       outline: none;
#       border-color: #38bdf8;
#       box-shadow: 0 0 0 1px #38bdf8;
#     }
#     .btn-primary {
#       background: linear-gradient(135deg, #38bdf8, #22c55e);
#       border: none;
#       font-weight: 600;
#       letter-spacing: 0.03em;
#     }
#     .badge-decision {
#       font-size: 0.8rem;
#       padding: 0.45rem 0.8rem;
#       border-radius: 999px;
#     }
#     .badge-apply {
#       background: rgba(34, 197, 94, 0.15);
#       color: #4ade80;
#       border: 1px solid rgba(74, 222, 128, 0.7);
#     }
#     .badge-maybe {
#       background: rgba(245, 158, 11, 0.15);
#       color: #fbbf24;
#       border: 1px solid rgba(251, 191, 36, 0.7);
#     }
#     .badge-reject {
#       background: rgba(248, 113, 113, 0.15);
#       color: #fca5a5;
#       border: 1px solid rgba(248, 113, 113, 0.7);
#     }
#     .form-label {
#       font-size: 0.85rem;
#       font-weight: 500;
#       color: #9ca3af;
#       text-transform: uppercase;
#       letter-spacing: 0.06em;
#     }
#     .footer-text {
#       font-size: 0.75rem;
#       color: #6b7280;
#     }
#   </style>
# </head>
# <body>
# <nav class="navbar navbar-dark mb-4" style="background: transparent;">
#   <div class="container">
#     <span class="navbar-brand">
#       AI Job Match Assistant
#     </span>
#   </div>
# </nav>

# <div class="container pb-5">
#   <div class="row g-4">
#     <!-- Input panel -->
#     <div class="col-lg-5">
#       <div class="card h-100">
#         <div class="card-header">
#           Input
#         </div>
#         <div class="card-body">
#           {% with messages = get_flashed_messages() %}
#           {% if messages %}
#             <div class="alert alert-warning py-2 small">
#               {{ messages[0] }}
#             </div>
#           {% endif %}
#           {% endwith %}
#           <form method="post" enctype="multipart/form-data">
#             <div class="mb-3">
#               <label class="form-label">Job Description</label>
#               <textarea name="job_description"
#                         class="form-control"
#                         rows="10"
#                         placeholder="Paste the full job description here...">{{ request.form.job_description or "" }}</textarea>
#             </div>
#             <div class="mb-3">
#               <label class="form-label">Resume File (PDF / DOCX / TXT)</label>
#               <input class="form-control form-control-sm"
#                      type="file"
#                      name="resume_file"
#                      accept=".pdf,.docx,.txt">
#               <div class="form-text footer-text mt-1">
#                 If no file is selected, existing resume text (if any) will be ignored.
#               </div>
#             </div>
#             <div class="mb-3">
#               <label class="form-label">Optional Context</label>
#               <textarea name="extra_context"
#                         class="form-control"
#                         rows="3"
#                         placeholder="Anything the agent should know (e.g., target industry, constraints)…">{{ request.form.extra_context or "" }}</textarea>
#             </div>
#             <button type="submit" class="btn btn-primary w-100 mt-2">
#               Analyze Match
#             </button>
#           </form>
#         </div>
#       </div>
#     </div>

#     <!-- Results panel -->
#     <div class="col-lg-7">
#       {% if result %}
#       <div class="row g-4">
#         <div class="col-12">
#           <div class="card">
#             <div class="card-header">
#               Overview
#             </div>
#             <div class="card-body row align-items-center">
#               <div class="col-md-4 text-center mb-3 mb-md-0">
#                 {% set score = result.match_score %}
#                 {% if score >= 70 %}
#                   {% set score_class = "score-high" %}
#                 {% elif score >= 50 %}
#                   {% set score_class = "score-medium" %}
#                 {% else %}
#                   {% set score_class = "score-low" %}
#                 {% endif %}
#                 <div class="score-circle {{ score_class }}">
#                   {{ score }}
#                 </div>
#                 <div class="mt-2 small text-uppercase text-muted">
#                   Match Score / 100
#                 </div>
#               </div>
#               <div class="col-md-8">
#                 <h5 class="mb-1">
#                   {{ result.job_analysis.role if result.job_analysis.role is string else (result.job_analysis.role | join(", ")) }}
#                 </h5>
#                 <div class="mb-2">
#                   {% if result.decision == "APPLY" %}
#                     <span class="badge-decision badge-apply">APPLY</span>
#                   {% elif result.decision == "MAYBE" %}
#                     <span class="badge-decision badge-maybe">MAYBE</span>
#                   {% else %}
#                     <span class="badge-decision badge-reject">REJECT</span>
#                   {% endif %}
#                   <span class="ms-2 small text-muted">Confidence: {{ result.confidence }}</span>
#                 </div>
#                 <div class="small text-muted">
#                   {{ result.reason }}
#                 </div>
#               </div>
#             </div>
#           </div>
#         </div>

#         <!-- Job analysis & resume analysis -->
#         <div class="col-md-6">
#           <div class="card h-100">
#             <div class="card-header">Job Analysis</div>
#             <div class="card-body small">
#               <div class="mb-2">
#                 <strong>Experience:</strong>
#                 <span class="ms-1">{{ result.job_analysis.experience_required }}</span>
#               </div>
#               <div class="mb-2">
#                 <strong>Keywords:</strong><br>
#                 {% for kw in result.job_analysis.keywords %}
#                   <span class="pill">{{ kw }}</span>
#                 {% else %}
#                   <span class="text-muted">None detected</span>
#                 {% endfor %}
#               </div>
#               <div class="mb-2">
#                 <strong>Hidden / Implied:</strong><br>
#                 {% for h in result.job_analysis.hidden_requirements %}
#                   <span class="pill">{{ h }}</span>
#                 {% else %}
#                   <span class="text-muted">None detected</span>
#                 {% endfor %}
#               </div>
#               <div class="mb-2">
#                 <strong>Skills List:</strong>
#                 <ul class="mt-1">
#                   {% for s in result.job_analysis.skills %}
#                     <li>{{ s }}</li>
#                   {% else %}
#                     <li class="text-muted">No structured skills parsed.</li>
#                   {% endfor %}
#                 </ul>
#               </div>
#             </div>
#           </div>
#         </div>

#         <div class="col-md-6">
#           <div class="card h-100">
#             <div class="card-header">Resume Analysis</div>
#             <div class="card-body small">
#               <div class="mb-2">
#                 <strong>Strengths vs Job:</strong>
#                 <ul class="mt-1">
#                   {% for s in result.resume_analysis.strengths %}
#                     <li>{{ s }}</li>
#                   {% else %}
#                     <li class="text-muted">No direct overlaps detected.</li>
#                   {% endfor %}
#                 </ul>
#               </div>
#               <div class="mb-2">
#                 <strong>Missing Skills:</strong>
#                 <ul class="mt-1">
#                   {% for s in result.resume_analysis.missing_skills %}
#                     <li>{{ s }}</li>
#                   {% else %}
#                     <li class="text-muted">No missing skills identified or no skills in JD.</li>
#                   {% endfor %}
#                 </ul>
#               </div>
#               <div class="mb-2">
#                 <strong>Weak Areas:</strong>
#                 <ul class="mt-1">
#                   {% for s in result.resume_analysis.weak_areas %}
#                     <li>{{ s }}</li>
#                   {% else %}
#                     <li class="text-muted">None detected.</li>
#                   {% endfor %}
#                 </ul>
#               </div>
#             </div>
#           </div>
#         </div>

#         <!-- Optimized resume & cover letter -->
#         <div class="col-12">
#           <div class="card">
#             <div class="card-header">Optimized Resume (Text)</div>
#             <div class="card-body">
#               <textarea class="form-control" rows="10" readonly>{{ result.optimized_resume }}</textarea>
#             </div>
#           </div>
#         </div>

#         <div class="col-12">
#           <div class="card">
#             <div class="card-header">Cover Letter Draft</div>
#             <div class="card-body">
#               <textarea class="form-control" rows="10" readonly>{{ result.cover_letter }}</textarea>
#             </div>
#           </div>
#         </div>
#       </div>
#       {% else %}
#       <div class="card">
#         <div class="card-body text-center py-5">
#           <h4 class="mb-2">Paste a job description and upload your resume to begin</h4>
#           <p class="text-muted mb-0 small">
#             The assistant will analyze role fit, highlight strengths & gaps, and generate an optimized resume and tailored cover letter.
#           </p>
#         </div>
#       </div>
#       {% endif %}
#     </div>
#   </div>

#   <div class="mt-4 text-center footer-text">
#     Built with Flask & Bootstrap. Paste, upload, analyze – no more raw JSON in the terminal.
#   </div>
# </div>

# <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/js/bootstrap.bundle.min.js"></script>
# </body>
# </html>
# """

# @app.route("/", methods=["GET", "POST"])
# def index():
#     result = None

#     if request.method == "POST":
#         job_description = request.form.get("job_description", "").strip()
#         extra_context = request.form.get("extra_context", "").strip()
#         resume_file = request.files.get("resume_file")

#         if not job_description:
#             flash("Please paste a job description.")
#             return render_template_string(TEMPLATE, result=None)

#         resume_text = ""
#         if resume_file and resume_file.filename:
#             filename = secure_filename(resume_file.filename)
#             save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
#             resume_file.save(save_path)
#             try:
#                 resume_text = extract_resume_text(save_path)
#             except Exception as e:
#                 flash(f"Could not read resume file: {e}")
#                 resume_text = ""
#         else:
#             flash("No resume file selected. Analysis will only use the job description.")

#         result = analyze_application(job_description, resume_text, extra_context)

#         # Wrap into a simple object-like thing so Jinja can use dot notation
#         result = json.loads(json.dumps(result, ensure_ascii=False), object_hook=lambda d: SimpleNamespace(**d))

#     from types import SimpleNamespace
#     return render_template_string(TEMPLATE, result=result)

# if __name__ == "__main__":
#     app.run(debug=True)


























import re
from collections import Counter
from typing import List, Dict, Any, Optional

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

# ----------------------------
# Constants
# ----------------------------

STOPWORDS = {
    "the", "and", "for", "with", "you", "your", "a", "an", "to", "of", "in",
    "on", "at", "by", "as", "is", "are", "be", "this", "that", "will", "our",
    "we", "they", "their", "from", "or", "it", "have", "has", "had", "i",
    "about", "us", "who", "what", "when", "where", "how", "why"
}

# ----------------------------
# Utility
# ----------------------------

def tokenize(text: str) -> List[str]:
    text = text.lower()
    text = re.sub(r"[^a-z0-9+]", " ", text)
    return [t for t in text.split() if t and t not in STOPWORDS]

# ----------------------------
# Resume text extraction
# ----------------------------

def extract_text_from_pdf(file_path: str) -> str:
    if not pdfplumber:
        raise ImportError("pdfplumber library not installed. Install with 'pip install pdfplumber'")
    text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text.append(page.extract_text() or "")
    return "\n".join(text)

def extract_text_from_docx(file_path: str) -> str:
    if not docx:
        raise ImportError("python-docx library not installed. Install with 'pip install python-docx'")
    document = docx.Document(file_path)
    return "\n".join([p.text for p in document.paragraphs])

def extract_resume_text(file_path: str) -> str:
    """Extract text from PDF, DOCX, or TXT."""
    lower = file_path.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_path)
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

# ----------------------------
# Job analysis helpers
# ----------------------------

def extract_role(job_description: str) -> str:
    lines = [l.strip() for l in job_description.splitlines() if l.strip()]
    return lines[0][:120] if lines else "insufficient data"

def extract_experience_required(job_description: str) -> str:
    match = re.search(r"(\d+)\+?\s+years", job_description, flags=re.IGNORECASE)
    if match:
        return f"{match.group(1)}+ years"
    range_match = re.search(r"(\d+)\s*-\s*(\d+)\s+years", job_description, flags=re.IGNORECASE)
    if range_match:
        return f"{range_match.group(1)}-{range_match.group(2)} years"
    return "not specified"

def extract_skills_from_section(text: str) -> List[str]:
    skills: List[str] = []
    lines = text.splitlines()
    headers = ("requirements", "responsibilities", "qualifications", "skills", "what you will do")
    current_section = None

    for line in lines:
        stripped = line.strip()
        lower = stripped.lower()

        if any(h in lower for h in headers):
            current_section = "skills"
            continue

        if current_section == "skills":
            # new header ends section
            if re.match(r"^[A-Z][A-Za-z\s]{0,40}:?$", stripped) and len(stripped.split()) <= 6:
                current_section = None
                continue

            if stripped.startswith(("-", "*", "•", "·")):
                item = stripped.lstrip("-*•· ").strip()
                if item:
                    skills.append(item)
            elif stripped and len(stripped) <= 80:
                skills.append(stripped)

    seen = set()
    unique = []
    for s in skills:
        key = s.lower()
        if key not in seen:
            seen.add(key)
            unique.append(s)
    return unique

def extract_keywords(job_description: str, max_keywords: int = 15) -> List[str]:
    tokens = tokenize(job_description)
    counts = Counter(tokens)
    common = [w for w, _ in counts.most_common() if len(w) > 3]
    return common[:max_keywords]

def extract_hidden_requirements(job_description: str) -> List[str]:
    text_lower = job_description.lower()
    patterns = {
        "fast-paced": "fast-paced environment",
        "self-starter": "self-starter",
        "work independently": "work independently",
        "cross-functional": "cross-functional collaboration",
        "stakeholders": "stakeholder management",
        "deadline": "ability to meet deadlines",
        "ambiguity": "comfort with ambiguity",
        "ownership": "sense of ownership",
        "leadership": "leadership responsibilities",
        "manage team": "team management"
    }
    hidden = [phrase for key, phrase in patterns.items() if key in text_lower]
    return list(dict.fromkeys(hidden))

# ----------------------------
# Resume analysis helpers
# ----------------------------

def extract_resume_skills(resume: str) -> List[str]:
    skills: List[str] = []
    lines = resume.splitlines()
    in_skills_section = False

    for line in lines:
        stripped = line.strip()
        lower = stripped.lower()

        if "skills" in lower and len(lower) <= 40:
            in_skills_section = True
            continue

        if in_skills_section:
            if re.match(r"^[A-Z][A-Za-z\s]{0,40}:?$", stripped) and "skills" not in lower:
                in_skills_section = False
            else:
                if stripped.startswith(("-", "*", "•", "·")):
                    items = stripped.lstrip("-*•· ").split(",")
                    skills.extend([i.strip() for i in items if i.strip()])
                elif "," in stripped:
                    skills.extend([i.strip() for i in stripped.split(",") if i.strip()])
                elif stripped:
                    skills.append(stripped)
                continue

        if stripped.startswith(("-", "*", "•", "·")):
            skills.append(stripped.lstrip("-*•· ").strip())

    seen = set()
    unique = []
    for s in skills:
        key = s.lower()
        if key not in seen:
            seen.add(key)
            unique.append(s)
    return unique

def get_overlap_keywords(job_keywords: List[str], resume_text: str) -> List[str]:
    resume_tokens = set(tokenize(resume_text))
    return [k for k in job_keywords if k.lower() in resume_tokens]

def analyze_resume_strengths(
    job_skills: List[str],
    resume_skills: List[str],
    resume_text: str
) -> Dict[str, List[str]]:
    strengths: List[str] = []
    missing: List[str] = []
    resume_lower = [s.lower() for s in resume_skills]

    for js in job_skills:
        if js.lower() in resume_lower:
            strengths.append(js)
        else:
            missing.append(js)

    weak_areas: List[str] = []
    if not re.search(r"\d+\+?\s+years", resume_text, flags=re.IGNORECASE):
        weak_areas.append("Years of experience not clearly specified")

    if job_skills and len(strengths) / len(job_skills) < 0.5:
        weak_areas.append("Significant gap between job-required skills and listed skills")

    return {"strengths": strengths, "missing": missing, "weak_areas": weak_areas}

# ----------------------------
# Scoring & decision
# ----------------------------

def compute_match_score(
    job_skills: List[str],
    resume_skills: List[str],
    job_keywords: List[str],
    resume_text: str
) -> int:
    if job_skills:
        skill_overlap = sum(1 for js in job_skills if js.lower() in [s.lower() for s in resume_skills])
        skill_score = (skill_overlap / len(job_skills)) * 70.0
    else:
        skill_score = 0.0

    if job_keywords:
        keyword_overlap = len(get_overlap_keywords(job_keywords, resume_text))
        keyword_score = (keyword_overlap / len(job_keywords)) * 30.0
    else:
        keyword_score = 0.0

    return max(0, min(100, int(round(skill_score + keyword_score))))

def decide_application(match_score: int) -> str:
    if match_score >= 70:
        return "APPLY"
    if match_score >= 50:
        return "MAYBE"
    return "REJECT"

def confidence_level(match_score: int, job_skills: List[str], resume_skills: List[str]) -> str:
    if not job_skills:
        return "LOW"
    overlap = sum(1 for js in job_skills if js.lower() in [s.lower() for s in resume_skills])
    if match_score >= 70 and overlap >= max(1, len(job_skills) // 2):
        return "HIGH"
    if 50 <= match_score < 70 and overlap > 0:
        return "MEDIUM"
    return "LOW"

# ----------------------------
# Resume optimization & cover letter
# ----------------------------

def build_optimized_resume(candidate_resume: str, job_keywords: List[str], resume_skills: List[str]) -> str:
    if not candidate_resume.strip():
        return "insufficient data"

    resume_tokens = set(tokenize(candidate_resume))
    kept_keywords = [k for k in job_keywords if k.lower() in resume_tokens]
    core_skills = [s for s in resume_skills if len(s) <= 60][:15]

    summary_lines: List[str] = []

    summary_lines.append("SUMMARY")
    if core_skills:
        summary_lines.append(
            "Experienced professional with strengths in: " + ", ".join(core_skills) + "."
        )
    else:
        summary_lines.append("Experienced professional seeking opportunities aligned with the attached job description.")

    if kept_keywords:
        summary_lines.append("")
        summary_lines.append("SKILLS & KEYWORDS")
        summary_lines.append(", ".join(sorted(set(kept_keywords))))

    # Do NOT invent new experience – just prepend summary to existing text
    return "\n".join(summary_lines) + "\n\nEXPERIENCE\n---------\n" + candidate_resume.strip()

def build_cover_letter(
    job_role: str,
    job_keywords: List[str],
    strengths: List[str],
    candidate_resume: str
) -> str:
    if not candidate_resume.strip():
        return "insufficient data"

    role = job_role if job_role != "insufficient data" else "the role"
    key_strengths = strengths[:5]
    resume_tokens = set(tokenize(candidate_resume))
    aligned_keywords = [k for k in job_keywords if k.lower() in resume_tokens][:5]

    strengths_text = ", ".join(key_strengths) if key_strengths else "the skills outlined in my resume"
    keywords_text = ", ".join(aligned_keywords) if aligned_keywords else ""

    p1 = (
        f"I am writing to express my interest in {role}. "
        f"After reviewing the job description, I believe my background and skills are well aligned with your needs."
    )
    if keywords_text:
        p2 = (
            f"Throughout my experience, I have developed strong capabilities in {keywords_text}. "
            f"These areas are reflected in my work, where I have applied them to deliver meaningful results."
        )
    else:
        p2 = (
            "My experience spans multiple responsibilities where I have consistently focused on delivering "
            "high-quality outcomes and collaborating effectively with stakeholders."
        )
    p3 = (
        f"My key strengths include {strengths_text}. "
        "I am comfortable taking ownership of my work, communicating clearly, and adapting to changing priorities."
    )
    p4 = (
        "I would welcome the opportunity to further discuss how my experience can contribute to your team’s goals. "
        "Thank you for considering my application."
    )
    return "\n\n".join([p1, p2, p3, p4])

# ----------------------------
# Job search strategy & interview plan
# ----------------------------

def build_job_search_strategy(
    job_description: str,
    candidate_resume_text: str,
    extra_context: Optional[str]
) -> Dict[str, Any]:
    role = extract_role(job_description) if job_description.strip() else "software jobs"
    experience_required = extract_experience_required(job_description) if job_description.strip() else "not specified"

    if experience_required != "not specified":
        base_query = f"{role} {experience_required}"
    else:
        base_query = role

    base_query = base_query.strip()

    platforms = [
        "LinkedIn Jobs",
        "Indeed",
        "Naukri",
        "Wellfound (AngelList)",
        "Company career pages"
    ]

    search_queries: List[str] = []
    if base_query:
        search_queries.append(base_query)
        search_queries.append(base_query + " remote")
        search_queries.append(base_query + " entry level" if "years" not in experience_required.lower() else base_query + " " + experience_required)

    tips = [
        "Use filters for location, experience level, and posted date to narrow down relevant roles.",
        "Save searches and turn on job alerts so new matching roles come to your inbox.",
        "For each APPLY decision, tailor your resume title and summary to the exact role wording.",
        "Prefer quality over quantity: for high-match roles, customize the resume and cover letter before applying."
    ]

    return {
        "platforms": platforms,
        "search_queries": search_queries,
        "tips": tips
    }

def build_interview_improvement_plan(
    job_analysis: Dict[str, Any],
    resume_analysis: Dict[str, Any]
) -> List[str]:
    plan: List[str] = []

    missing_skills = resume_analysis.get("missing_skills", []) or []
    weak_areas = resume_analysis.get("weak_areas", []) or []

    if missing_skills:
        plan.append(
            "Review the job-required skills and explicitly mention any of these skills you truly have experience with in your resume bullets."
        )
        for s in missing_skills[:5]:
            plan.append(
                f"If you have real experience with '{s}', highlight it clearly in your summary, skills, or project descriptions."
            )

    for w in weak_areas:
        if "Years of experience not clearly specified" in w:
            plan.append(
                "Add clear years-of-experience statements for your main roles (e.g., '2+ years Python development experience')."
            )
        else:
            plan.append(f"Address this resume weakness: {w}")

    if not plan:
        plan.append(
            "Your resume is reasonably aligned. Focus on tailoring the top section (SUMMARY and SKILLS) to each high-match job."
        )

    return plan

# ----------------------------
# Main analysis — returns STRICT JSON spec
# ----------------------------

def analyze_application(
    job_description: str,
    candidate_resume_text: str,
    extra_context: Optional[str] = None
) -> Dict[str, Any]:
    job_description = job_description or ""
    candidate_resume_text = candidate_resume_text or ""

    # Job analysis
    role = extract_role(job_description) if job_description.strip() else "insufficient data"
    experience_required = extract_experience_required(job_description) if job_description.strip() else "insufficient data"
    job_skills = extract_skills_from_section(job_description) if job_description.strip() else []
    keywords = extract_keywords(job_description) if job_description.strip() else []
    hidden = extract_hidden_requirements(job_description) if job_description.strip() else []

    job_analysis = {
        "role": role,
        "skills": job_skills,
        "experience_required": experience_required,
        "keywords": keywords or ["insufficient data"],
        "hidden_requirements": hidden or ["insufficient data"]
    }

    # Resume analysis
    resume_skills = extract_resume_skills(candidate_resume_text) if candidate_resume_text.strip() else []
    if candidate_resume_text.strip():
        strength_analysis = analyze_resume_strengths(job_skills, resume_skills, candidate_resume_text)
        resume_strengths = strength_analysis["strengths"]
        missing_skills = strength_analysis["missing"]
        weak_areas = strength_analysis["weak_areas"]
    else:
        resume_strengths = []
        missing_skills = job_skills
        weak_areas = ["Candidate resume not provided or empty"]

    resume_analysis = {
        "strengths": resume_strengths,
        "missing_skills": missing_skills,
        "weak_areas": weak_areas
    }

    # Scoring, decision, confidence
    match_score = compute_match_score(job_skills, resume_skills, keywords, candidate_resume_text)
    decision = decide_application(match_score)
    confidence = confidence_level(match_score, job_skills, resume_skills)

    # Reason text
    if decision == "APPLY":
        reason = f"Match score {match_score}: substantial overlap; suitable to submit application."
    elif decision == "MAYBE":
        reason = f"Match score {match_score}: partial overlap; may be worthwhile."
    else:
        reason = f"Match score {match_score}: limited overlap; unlikely to be competitive."

    # Optimized resume & cover letter
    optimized_resume = build_optimized_resume(candidate_resume_text, keywords, resume_skills)
    cover_letter = build_cover_letter(role, keywords, resume_strengths, candidate_resume_text)

    # Job search strategy & interview improvement
    job_search_strategy = build_job_search_strategy(job_description, candidate_resume_text, extra_context)
    interview_improvement_plan = build_interview_improvement_plan(job_analysis, resume_analysis)

    # STRICT JSON spec
    result: Dict[str, Any] = {
        "job_analysis": job_analysis,
        "resume_analysis": resume_analysis,
        "match_score": match_score,
        "decision": decision,
        "confidence": confidence,
        "optimized_resume": optimized_resume,
        "cover_letter": cover_letter,
        "job_search_strategy": job_search_strategy,
        "interview_improvement_plan": interview_improvement_plan
    }
    return result